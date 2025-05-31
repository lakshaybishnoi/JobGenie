import os
import re
import PyPDF2
import docx
from typing import Dict, List, Any
from pathlib import Path

class ResumeParser:
    """Resume parser for extracting information from PDF, DOCX, and TXT files."""
    
    def __init__(self):
        # Common skill patterns
        self.skill_patterns = [
            # Programming languages
            r'\b(?:python|java|javascript|typescript|c\+\+|c#|php|ruby|go|rust|swift|kotlin|scala|r|matlab|sql)\b',
            # Web technologies
            r'\b(?:html|css|react|angular|vue|django|flask|spring|laravel|express|node\.js|jquery|bootstrap)\b',
            # Databases
            r'\b(?:mysql|postgresql|mongodb|redis|oracle|sqlserver|sqlite|elasticsearch|cassandra)\b',
            # Cloud & DevOps
            r'\b(?:aws|azure|gcp|docker|kubernetes|jenkins|git|ci/cd|devops|terraform|ansible)\b',
            # Data Science & ML
            r'\b(?:pandas|numpy|scikit-learn|tensorflow|pytorch|jupyter|tableau|power bi|excel|spss)\b',
            # Tools & Frameworks
            r'\b(?:rest|api|graphql|microservices|agile|scrum|linux|windows|jira|confluence)\b'
        ]
        
        # Education patterns
        self.education_patterns = [
            r'(?:bachelor|b\.?s\.?|b\.?a\.?|undergraduate)[\s\w]*(?:in|of)?\s*[\w\s]+',
            r'(?:master|m\.?s\.?|m\.?a\.?|mba|graduate)[\s\w]*(?:in|of)?\s*[\w\s]+',
            r'(?:phd|ph\.?d\.?|doctorate|doctoral)[\s\w]*(?:in|of)?\s*[\w\s]+',
            r'(?:associate|diploma|certificate)[\s\w]*(?:in|of)?\s*[\w\s]+',
            r'(?:computer science|engineering|mathematics|physics|chemistry|biology|business|economics|finance|marketing)'
        ]
        
        # Experience keywords
        self.experience_keywords = [
            'experience', 'worked', 'developed', 'managed', 'led', 'built', 'created', 
            'designed', 'implemented', 'maintained', 'deployed', 'optimized', 'analyzed',
            'collaborated', 'coordinated', 'supervised', 'mentored', 'trained', 'architected'
        ]
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse resume and extract relevant information."""
        file_extension = Path(file_path).suffix.lower()
        
        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self._extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text = self._extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if not text.strip():
                raise ValueError("No text content found in the resume")
            
            # Extract information
            skills = self._extract_skills(text)
            education = self._extract_education(text)
            experience = self._extract_experience(text)
            contact_info = self._extract_contact_info(text)
            
            return {
                'filename': os.path.basename(file_path),
                'text': text,
                'skills': skills,
                'education': education,
                'experience': experience,
                'contact_info': contact_info,
                'word_count': len(text.split()),
                'parsing_status': 'success'
            }
            
        except Exception as e:
            return {
                'filename': os.path.basename(file_path),
                'text': '',
                'skills': [],
                'education': [],
                'experience': [],
                'contact_info': {},
                'word_count': 0,
                'parsing_status': f'error: {str(e)}'
            }
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
        
        return text.strip()
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                raise ValueError(f"Error reading TXT file: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error reading TXT file: {str(e)}")
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        skills = set()
        text_lower = text.lower()
        
        # Use predefined skill patterns
        for pattern in self.skill_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                skill = match.group().strip()
                if len(skill) > 1:  # Avoid single characters
                    skills.add(skill.title())
        
        # Look for skills in common sections
        skill_sections = self._find_sections(text, ['skills', 'technical skills', 'technologies', 'tools'])
        for section in skill_sections:
            # Extract comma-separated or bullet-pointed skills
            section_lower = section.lower()
            
            # Split by common delimiters
            potential_skills = re.split(r'[,•\n\r\t\|/]', section_lower)
            
            for skill in potential_skills:
                skill = skill.strip()
                # Check if it matches known patterns or is a reasonable skill name
                if len(skill) > 2 and len(skill) < 30:
                    # Remove common non-skill words
                    if not re.match(r'^(?:and|or|with|using|including|such as|etc|experience|years?)$', skill):
                        skills.add(skill.title())
        
        return sorted(list(skills))
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information from resume text."""
        education = set()
        text_lower = text.lower()
        
        # Use predefined education patterns
        for pattern in self.education_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                edu = match.group().strip()
                if len(edu) > 5:  # Reasonable length for education entry
                    education.add(edu.title())
        
        # Look for education in common sections
        edu_sections = self._find_sections(text, ['education', 'academic', 'qualification', 'degree'])
        for section in edu_sections:
            # Extract lines that likely contain education info
            lines = section.split('\n')
            for line in lines:
                line = line.strip()
                if len(line) > 10 and any(keyword in line.lower() for keyword in 
                                        ['university', 'college', 'institute', 'school', 'degree', 'bachelor', 'master', 'phd']):
                    education.add(line)
        
        return sorted(list(education))
    
    def _extract_experience(self, text: str) -> List[str]:
        """Extract work experience information from resume text."""
        experience = []
        
        # Look for experience sections
        exp_sections = self._find_sections(text, ['experience', 'work experience', 'employment', 'professional experience', 'career'])
        
        for section in exp_sections:
            # Split by potential job entries (often separated by company names or dates)
            lines = section.split('\n')
            current_job = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_job:
                        experience.append(' '.join(current_job))
                        current_job = []
                    continue
                
                # Check if this line starts a new job (often has dates or company indicators)
                if re.search(r'\b(?:20\d{2}|19\d{2})\b', line) or any(keyword in line.lower() for keyword in ['company', 'corp', 'inc', 'ltd']):
                    if current_job:
                        experience.append(' '.join(current_job))
                        current_job = []
                
                current_job.append(line)
            
            # Add the last job if exists
            if current_job:
                experience.append(' '.join(current_job))
        
        return experience[:10]  # Limit to top 10 experience entries
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume text."""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern
        phone_pattern = r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([A-Za-z0-9_-]+)'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
        
        # GitHub pattern
        github_pattern = r'(?:github\.com/)([A-Za-z0-9_-]+)'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = f"github.com/{github_match.group(1)}"
        
        return contact_info
    
    def _find_sections(self, text: str, section_keywords: List[str]) -> List[str]:
        """Find sections in the resume based on keywords."""
        sections = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            if any(keyword in line_lower for keyword in section_keywords):
                # Extract the section content
                section_content = []
                j = i + 1
                
                # Continue until we find another section or end of text
                while j < len(lines):
                    next_line = lines[j].strip()
                    
                    # Stop if we hit another section header
                    if (next_line and 
                        any(header in next_line.lower() for header in 
                            ['education', 'experience', 'skills', 'objective', 'summary', 'projects', 'certifications', 'awards'])):
                        break
                    
                    section_content.append(next_line)
                    j += 1
                
                sections.append('\n'.join(section_content))
        
        return sections
    
    def validate_resume_content(self, resume_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and score the quality of extracted resume content."""
        score = 0
        feedback = []
        
        # Check for skills
        if resume_data.get('skills'):
            score += 30
            feedback.append(f"✅ Found {len(resume_data['skills'])} skills")
        else:
            feedback.append("❌ No technical skills detected")
        
        # Check for education
        if resume_data.get('education'):
            score += 25
            feedback.append(f"✅ Found {len(resume_data['education'])} education entries")
        else:
            feedback.append("❌ No education information detected")
        
        # Check for experience
        if resume_data.get('experience'):
            score += 30
            feedback.append(f"✅ Found {len(resume_data['experience'])} experience entries")
        else:
            feedback.append("❌ No work experience detected")
        
        # Check for contact info
        contact_score = len(resume_data.get('contact_info', {})) * 3.75  # Up to 15 points
        score += contact_score
        if contact_score > 0:
            feedback.append(f"✅ Found contact information")
        else:
            feedback.append("❌ No contact information detected")
        
        # Check text length
        word_count = resume_data.get('word_count', 0)
        if word_count > 100:
            feedback.append(f"✅ Good content length ({word_count} words)")
        else:
            feedback.append("⚠️ Resume content seems too short")
        
        return {
            'quality_score': min(100, score),
            'feedback': feedback,
            'recommendations': self._get_recommendations(resume_data)
        }
    
    def _get_recommendations(self, resume_data: Dict[str, Any]) -> List[str]:
        """Get recommendations for improving resume content."""
        recommendations = []
        
        if not resume_data.get('skills'):
            recommendations.append("Add a clear skills section with technical competencies")
        
        if not resume_data.get('education'):
            recommendations.append("Include education background and qualifications")
        
        if not resume_data.get('experience'):
            recommendations.append("Add work experience with specific achievements")
        
        if not resume_data.get('contact_info', {}).get('email'):
            recommendations.append("Include contact email address")
        
        if len(resume_data.get('skills', [])) < 5:
            recommendations.append("Consider adding more relevant technical skills")
        
        return recommendations
